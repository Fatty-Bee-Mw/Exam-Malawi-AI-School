"""
File Processing Utilities for Training System
Supports multiple file formats for educational content extraction
"""

import os
import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional
from io import BytesIO
import re

# Optional imports for different file formats
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import xml.etree.ElementTree as ET
    XML_AVAILABLE = True
except ImportError:
    XML_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

import csv
import zipfile

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles extraction of text content from various file formats"""
    
    SUPPORTED_FORMATS = {
        '.txt': 'Plain Text',
        '.pdf': 'PDF Document',
        '.docx': 'Word Document',
        '.doc': 'Word Document (Legacy)',
        '.csv': 'CSV Data',
        '.xlsx': 'Excel Spreadsheet',
        '.xls': 'Excel Spreadsheet (Legacy)',
        '.json': 'JSON Data',
        '.xml': 'XML Document',
        '.html': 'HTML Document',
        '.htm': 'HTML Document',
        '.md': 'Markdown',
        '.rtf': 'Rich Text Format',
        '.odt': 'OpenDocument Text',
        '.ods': 'OpenDocument Spreadsheet'
    }
    
    def __init__(self):
        self.check_dependencies()
    
    def check_dependencies(self):
        """Check which file processing libraries are available"""
        missing = []
        if not PDF_AVAILABLE:
            missing.append("PyPDF2 (for PDF files)")
        if not DOCX_AVAILABLE:
            missing.append("python-docx (for Word documents)")
        if not PANDAS_AVAILABLE:
            missing.append("pandas (for Excel/CSV files)")
        if not BS4_AVAILABLE:
            missing.append("beautifulsoup4 (for HTML files)")
        
        if missing:
            logger.warning(f"Some file processing libraries are missing: {', '.join(missing)}")
    
    def extract_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from file content based on file extension
        Returns dict with extracted text and metadata
        """
        file_ext = Path(filename).suffix.lower()
        
        result = {
            "text": "",
            "format": file_ext,
            "format_name": self.SUPPORTED_FORMATS.get(file_ext, "Unknown"),
            "success": False,
            "error": None,
            "metadata": {}
        }
        
        try:
            if file_ext == '.txt':
                result["text"] = self._extract_from_txt(file_content)
            elif file_ext == '.pdf':
                result["text"] = self._extract_from_pdf(file_content)
            elif file_ext in ['.doc', '.docx']:
                result["text"] = self._extract_from_docx(file_content)
            elif file_ext == '.csv':
                result["text"] = self._extract_from_csv(file_content)
            elif file_ext in ['.xlsx', '.xls']:
                result["text"] = self._extract_from_excel(file_content)
            elif file_ext == '.json':
                result["text"] = self._extract_from_json(file_content)
            elif file_ext == '.xml':
                result["text"] = self._extract_from_xml(file_content)
            elif file_ext in ['.html', '.htm']:
                result["text"] = self._extract_from_html(file_content)
            elif file_ext == '.md':
                result["text"] = self._extract_from_markdown(file_content)
            elif file_ext == '.rtf':
                result["text"] = self._extract_from_rtf(file_content)
            else:
                # Try to decode as plain text for unknown formats
                result["text"] = self._extract_from_txt(file_content)
            
            # Clean and validate extracted text
            result["text"] = self._clean_text(result["text"])
            
            if result["text"]:
                result["success"] = True
                result["metadata"]["word_count"] = len(result["text"].split())
                result["metadata"]["char_count"] = len(result["text"])
            else:
                result["error"] = "No text content extracted"
                
        except Exception as e:
            result["error"] = str(e)
            logger.error(f"Error extracting text from {filename}: {e}")
        
        return result
    
    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from plain text files"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Fallback with error handling
        return file_content.decode('utf-8', errors='ignore')
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF files"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available for PDF processing")
        
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
            
            return text
        except Exception as e:
            raise Exception(f"PDF extraction failed: {e}")
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for Word document processing")
        
        try:
            doc_file = BytesIO(file_content)
            doc = docx.Document(doc_file)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {e}")
    
    def _extract_from_csv(self, file_content: bytes) -> str:
        """Extract text from CSV files"""
        try:
            csv_text = file_content.decode('utf-8', errors='ignore')
            csv_reader = csv.reader(csv_text.splitlines())
            text = ""
            
            for row_num, row in enumerate(csv_reader):
                if row:  # Skip empty rows
                    if row_num == 0:
                        # Header row
                        text += "Headers: " + " | ".join(row) + "\n\n"
                    else:
                        text += f"Row {row_num}: " + " | ".join(row) + "\n"
            
            return text
        except Exception as e:
            raise Exception(f"CSV extraction failed: {e}")
    
    def _extract_from_excel(self, file_content: bytes) -> str:
        """Extract text from Excel files"""
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas not available for Excel processing")
        
        try:
            excel_file = BytesIO(file_content)
            # Read all sheets
            sheets = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
            text = ""
            
            for sheet_name, df in sheets.items():
                text += f"\n=== Sheet: {sheet_name} ===\n"
                
                # Convert DataFrame to readable text
                if not df.empty:
                    # Add column headers
                    text += "Columns: " + " | ".join(df.columns.astype(str)) + "\n\n"
                    
                    # Add data rows (limit to prevent huge files)
                    for idx, row in df.head(100).iterrows():  # Limit to first 100 rows
                        row_text = " | ".join(row.astype(str))
                        text += f"Row {idx + 1}: {row_text}\n"
                    
                    if len(df) > 100:
                        text += f"\n... and {len(df) - 100} more rows\n"
                else:
                    text += "Empty sheet\n"
                
                text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Excel extraction failed: {e}")
    
    def _extract_from_json(self, file_content: bytes) -> str:
        """Extract text from JSON files"""
        try:
            json_text = file_content.decode('utf-8', errors='ignore')
            json_data = json.loads(json_text)
            return self._json_to_text(json_data)
        except Exception as e:
            raise Exception(f"JSON extraction failed: {e}")
    
    def _json_to_text(self, data, prefix="", max_depth=10) -> str:
        """Convert JSON data to readable text"""
        if max_depth <= 0:
            return f"{prefix}[Max depth reached]\n"
        
        text = ""
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    text += f"{prefix}{key}:\n"
                    text += self._json_to_text(value, prefix + "  ", max_depth - 1)
                else:
                    text += f"{prefix}{key}: {value}\n"
        elif isinstance(data, list):
            for i, item in enumerate(data[:50]):  # Limit to first 50 items
                text += f"{prefix}Item {i+1}:\n"
                text += self._json_to_text(item, prefix + "  ", max_depth - 1)
            if len(data) > 50:
                text += f"{prefix}... and {len(data) - 50} more items\n"
        else:
            text += f"{prefix}{data}\n"
        return text
    
    def _extract_from_xml(self, file_content: bytes) -> str:
        """Extract text from XML files"""
        if not XML_AVAILABLE:
            raise ImportError("xml.etree.ElementTree not available")
        
        try:
            xml_text = file_content.decode('utf-8', errors='ignore')
            root = ET.fromstring(xml_text)
            return self._xml_to_text(root)
        except Exception as e:
            raise Exception(f"XML extraction failed: {e}")
    
    def _xml_to_text(self, element, prefix="", max_depth=10) -> str:
        """Convert XML element to readable text"""
        if max_depth <= 0:
            return f"{prefix}[Max depth reached]\n"
        
        text = ""
        
        # Add element text content
        if element.text and element.text.strip():
            text += f"{prefix}{element.tag}: {element.text.strip()}\n"
        else:
            text += f"{prefix}{element.tag}:\n"
        
        # Add attributes
        if element.attrib:
            for attr, value in element.attrib.items():
                text += f"{prefix}  @{attr}: {value}\n"
        
        # Process children
        for child in element:
            text += self._xml_to_text(child, prefix + "  ", max_depth - 1)
        
        return text
    
    def _extract_from_html(self, file_content: bytes) -> str:
        """Extract text from HTML files"""
        try:
            html_text = file_content.decode('utf-8', errors='ignore')
            
            if BS4_AVAILABLE:
                soup = BeautifulSoup(html_text, 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                return soup.get_text()
            else:
                # Fallback without BeautifulSoup
                # Remove HTML tags
                clean_text = re.sub('<.*?>', '', html_text)
                # Decode HTML entities
                clean_text = clean_text.replace('&nbsp;', ' ')
                clean_text = clean_text.replace('&amp;', '&')
                clean_text = clean_text.replace('&lt;', '<')
                clean_text = clean_text.replace('&gt;', '>')
                return clean_text
        except Exception as e:
            raise Exception(f"HTML extraction failed: {e}")
    
    def _extract_from_markdown(self, file_content: bytes) -> str:
        """Extract text from Markdown files"""
        try:
            md_text = file_content.decode('utf-8', errors='ignore')
            
            # Remove markdown formatting while preserving content
            # Remove headers
            md_text = re.sub(r'^#+\s*', '', md_text, flags=re.MULTILINE)
            # Remove bold/italic
            md_text = re.sub(r'\*\*(.*?)\*\*', r'\1', md_text)
            md_text = re.sub(r'\*(.*?)\*', r'\1', md_text)
            # Remove links but keep text
            md_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', md_text)
            # Remove code blocks
            md_text = re.sub(r'```.*?```', '', md_text, flags=re.DOTALL)
            md_text = re.sub(r'`([^`]+)`', r'\1', md_text)
            
            return md_text
        except Exception as e:
            raise Exception(f"Markdown extraction failed: {e}")
    
    def _extract_from_rtf(self, file_content: bytes) -> str:
        """Extract text from RTF files"""
        try:
            rtf_text = file_content.decode('utf-8', errors='ignore')
            
            # Basic RTF text extraction (remove RTF commands)
            # Remove RTF control words
            clean_text = re.sub(r'\\[a-z]+\d*\s?', '', rtf_text)
            # Remove remaining RTF syntax
            clean_text = re.sub(r'[{}]', '', clean_text)
            # Clean up whitespace
            clean_text = re.sub(r'\s+', ' ', clean_text)
            
            return clean_text.strip()
        except Exception as e:
            raise Exception(f"RTF extraction failed: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize paragraph breaks
        
        # Remove special characters but keep educational content
        # Keep letters, numbers, basic punctuation, and educational symbols
        text = re.sub(r'[^\w\s\.\,\?\!\:\;\-\(\)\[\]\"\'\/\+\=\%\$\#\@\&\*]', ' ', text)
        
        # Remove extra spaces
        text = ' '.join(text.split())
        
        # Ensure minimum length for educational content
        if len(text.strip()) < 20:
            return ""
        
        return text.strip()
    
    def get_supported_formats(self) -> Dict[str, str]:
        """Get list of supported file formats"""
        return self.SUPPORTED_FORMATS.copy()
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported"""
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.SUPPORTED_FORMATS

# Global file processor instance
file_processor = FileProcessor()
